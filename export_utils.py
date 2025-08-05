from docx import Document
from PyQt6.QtWidgets import QFileDialog, QMessageBox, QDialog, QVBoxLayout, QLabel, QLineEdit, QPushButton, QFormLayout
from PyQt6.QtGui import QFont
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import datetime
import os


class HeaderFooterDialog(QDialog):
    """Custom dialog for setting advanced header/footer with all fields in one GUI, with larger input boxes."""

    def __init__(self, parent, title, current_data):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.current_data = current_data
        self.setMinimumSize(600, 400)  # Larger dialog size
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        # Form for fields with larger font and wider input boxes
        form_layout = QFormLayout()
        font = QFont("Arial", 12)  # Larger font for labels and inputs

        self.name_edit = QLineEdit(self.current_data.get('name', ''))
        self.name_edit.setFont(font)
        self.name_edit.setMinimumWidth(400)  # Wider input box
        form_layout.addRow(QLabel("Name:", font=font), self.name_edit)

        self.church_edit = QLineEdit(self.current_data.get('church', ''))
        self.church_edit.setFont(font)
        self.church_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Church:", font=font), self.church_edit)

        self.organization_edit = QLineEdit(self.current_data.get('organization', ''))
        self.organization_edit.setFont(font)
        self.organization_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Organization:", font=font), self.organization_edit)

        self.email_edit = QLineEdit(self.current_data.get('email', ''))
        self.email_edit.setFont(font)
        self.email_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Email:", font=font), self.email_edit)

        self.phone_edit = QLineEdit(self.current_data.get('phone', ''))
        self.phone_edit.setFont(font)
        self.phone_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Phone:", font=font), self.phone_edit)

        self.website_edit = QLineEdit(self.current_data.get('website', ''))
        self.website_edit.setFont(font)
        self.website_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Website:", font=font), self.website_edit)

        self.additional_edit = QLineEdit(self.current_data.get('additional', ''))
        self.additional_edit.setFont(font)
        self.additional_edit.setMinimumWidth(400)
        form_layout.addRow(QLabel("Additional Info:", font=font), self.additional_edit)

        self.layout.addLayout(form_layout)

        # Save button
        save_btn = QPushButton("Save")
        save_btn.setFont(font)
        save_btn.setStyleSheet("padding: 10px;")  # Larger button
        save_btn.clicked.connect(self.accept)
        self.layout.addWidget(save_btn)

    def get_data(self):
        """Return the updated data as a dictionary."""
        return {
            'name': self.name_edit.text().strip(),
            'church': self.church_edit.text().strip(),
            'organization': self.organization_edit.text().strip(),
            'email': self.email_edit.text().strip(),
            'phone': self.phone_edit.text().strip(),
            'website': self.website_edit.text().strip(),
            'additional': self.additional_edit.text().strip()
        }


def set_header(parent, sermon, status_bar):
    """Set advanced header with all fields in one GUI dialog."""
    try:
        if not isinstance(sermon.get('header'), dict):
            sermon['header'] = {}
        current_data = sermon['header']

        dialog = HeaderFooterDialog(parent, "Set Header", current_data)
        if dialog.exec():
            sermon['header'] = dialog.get_data()
            status_bar.showMessage("Header set.", 3000)
    except Exception as e:
        status_bar.showMessage(f"Error setting header: {str(e)}", 5000)


def set_footer(parent, sermon, status_bar):
    """Set advanced footer with all fields in one GUI dialog."""
    try:
        if not isinstance(sermon.get('footer'), dict):
            sermon['footer'] = {}
        current_data = sermon['footer']

        dialog = HeaderFooterDialog(parent, "Set Footer", current_data)
        if dialog.exec():
            sermon['footer'] = dialog.get_data()
            status_bar.showMessage("Footer set.", 3000)
    except Exception as e:
        status_bar.showMessage(f"Error setting footer: {str(e)}", 5000)


def save_as_word(parent, sermon, status_bar, filename=None):
    try:
        # Initialize last_save_dir if not set
        if not hasattr(parent, 'last_save_dir'):
            parent.last_save_dir = os.getcwd()  # Fallback to current working directory

        # Prompt user for save location if no filename provided
        if not filename:
            title = sermon.get('title', 'Untitled_Sermon').replace(' ', '_')
            date = datetime.now().strftime('%Y-%m-%d')
            folder_name = f"{title}_{date}"
            default_filename = f"{title}_{date}.docx"
            folder_path = os.path.join(parent.last_save_dir, folder_name)
            filename, _ = QFileDialog.getSaveFileName(
                parent,
                "Save Sermon As",
                os.path.join(folder_path, default_filename),
                "Word Documents (*.docx)"
            )
        if not filename:
            return

        # Ensure .docx extension
        if not filename.lower().endswith('.docx'):
            filename = os.path.splitext(filename)[0] + '.docx'

        # Create folder if it doesn't exist
        os.makedirs(os.path.dirname(filename), exist_ok=True)

        doc = Document()
        section = doc.sections[0]

        # Advanced Header: Only add non-empty fields
        header = section.header
        header_p = header.add_paragraph()
        header_data = sermon.get('header', {})
        header_fields = [
            f"Name: {header_data['name']}" if header_data.get('name') else None,
            f"Church: {header_data['church']}" if header_data.get('church') else None,
            f"Organization: {header_data['organization']}" if header_data.get('organization') else None,
            f"Email: {header_data['email']}" if header_data.get('email') else None,
            f"Phone: {header_data['phone']}" if header_data.get('phone') else None,
            f"Website: {header_data['website']}" if header_data.get('website') else None,
            f"Additional Info: {header_data['additional']}" if header_data.get('additional') else None
        ]
        header_text = '\n'.join(field for field in header_fields if field)
        if header_text:
            header_p.add_run(header_text)
            header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_p.runs:
                run.font.size = Pt(10)  # Smaller font for header

        # Advanced Footer: Only add non-empty fields
        footer = section.footer
        footer_p = footer.add_paragraph()
        footer_data = sermon.get('footer', {})
        footer_fields = [
            f"Name: {footer_data['name']}" if footer_data.get('name') else None,
            f"Church: {footer_data['church']}" if footer_data.get('church') else None,
            f"Organization: {footer_data['organization']}" if footer_data.get('organization') else None,
            f"Email: {footer_data['email']}" if footer_data.get('email') else None,
            f"Phone: {footer_data['phone']}" if footer_data.get('phone') else None,
            f"Website: {footer_data['website']}" if footer_data.get('website') else None,
            f"Additional Info: {footer_data['additional']}" if footer_data.get('additional') else None
        ]
        footer_text = '\n'.join(field for field in footer_fields if field)
        if footer_text:
            footer_p.add_run(footer_text)
            footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in footer_p.runs:
                run.font.size = Pt(10)  # Smaller font for footer

        # Title
        title = doc.add_heading(sermon.get('title', 'Untitled Sermon'), level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Intro
        doc.add_heading("Introduction", level=2)
        intro_text = sermon.get('intro', 'No introduction provided.')
        doc.add_paragraph(intro_text, style='Normal')

        # Sermon Content
        doc.add_heading("Sermon Content", level=2)
        content_text = sermon.get('content', 'No sermon content provided.')
        doc.add_paragraph(content_text, style='Normal')

        # Verses and Notes
        doc.add_heading("Verses and Notes", level=2)
        verses_notes = sermon.get('verses_notes', [])
        if not verses_notes:
            doc.add_paragraph("No verses or notes provided.", style='Normal')
        else:
            for vn in verses_notes:
                ref = vn.get('ref', 'Unknown')
                text = vn.get('text', '')
                note = vn.get('note', '')
                doc.add_paragraph(f"{ref}: {text}", style='Normal')
                if note:
                    doc.add_paragraph(f"Note: {note}", style='Normal')

        # Add spacing at the end
        doc.add_paragraph()

        # Save the document
        doc.save(filename)
        status_bar.showMessage(f"Exported to {filename}.", 3000)
        parent.last_save_dir = os.path.dirname(filename)
    except Exception as e:
        QMessageBox.warning(parent, "Export Error", f"Failed to export: {str(e)}")
